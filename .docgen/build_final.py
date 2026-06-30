# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TPL = "/mnt/c/Users/Tup20/OneDrive - Universitatea Politehnica Timisoara/Desktop/documentatii licenta/Template AC ro 2026.docx"
OUT = "/mnt/c/Users/Tup20/OneDrive - Universitatea Politehnica Timisoara/Desktop/documentatii licenta/Documentatie_Licenta_ResoLab.docx"
LOCAL = "/home/tup20/licenta/Documentatie_Licenta_ResoLab.docx"
FIGS = os.path.join(os.path.dirname(__file__), 'figs')

AUTHOR = "Tristan-Dumitru CVASNIUC"
AUTHOR_HDR = "Tristan-Dumitru Cvasniuc"
COORD = "Conf.dr.ing. Lucian PRODAN"
TITLE = "ResoLab – mini-DAW educațional pentru procesarea audio în timp real în browser folosind Rust și WebAssembly"
TITLE_SHORT = "ResoLab – mini-DAW educațional în browser"
PROGRAM = "Informatică"; YEAR = "2025–2026"

REZUMAT = ("Lucrarea de față prezintă ResoLab, o aplicație web pentru învățarea și exersarea prelucrării sunetului, "
"care funcționează integral în browser, fără instalare și fără un server de procesare. Spre deosebire de editoarele "
"audio obișnuite, în care efectele se aplică fără a explica ce se întâmplă cu semnalul, aplicația a fost concepută "
"astfel încât fiecare etapă a procesării să fie vizibilă și explicabilă. Nucleul de prelucrare este scris în limbajul "
"Rust și compilat către WebAssembly, ceea ce permite atingerea unei latențe reduse direct în pagina web; el rulează "
"izolat, pe firul audio al browserului, printr-un AudioWorklet care servește drept punte subțire către interfața "
"construită cu React și TypeScript. Aplicația oferă un lanț reconfigurabil de șaisprezece efecte, un sintetizator "
"polifonic, un modul de înregistrare cu selecția interfeței audio externe, un metronom cu programare temporală "
"precisă, vizualizări spectrale și un strat educațional care analizează configurația curentă și propune observații "
"și recomandări. Rezultatul poate fi exportat ca fișier WAV. Testele efectuate au confirmat că motorul procesează "
"lanțuri complexe fără întreruperi audibile, validând abordarea aleasă pentru un instrument educațional interactiv.")
ABSTRACT = ("This thesis presents ResoLab, a web application for learning and practising sound engineering that runs "
"entirely in the browser, with no installation and no processing server. Unlike conventional audio editors, where "
"effects are applied without explaining what happens to the signal, the application was designed so that every "
"processing step is visible and explainable. The processing core is written in Rust and compiled to WebAssembly, "
"which enables low-latency processing directly in the web page; it runs in isolation, on the browser audio thread, "
"through an AudioWorklet that acts as a thin bridge to the interface built with React and TypeScript. The application "
"provides a reconfigurable chain of sixteen effects, a polyphonic synthesiser, a recording module with external "
"audio-interface selection, a metronome with precise temporal scheduling, spectral visualisations and an educational "
"layer that analyses the current configuration and offers observations and recommendations. The result can be exported "
"as a WAV file. Testing confirmed that the engine processes complex chains without audible dropouts, validating the "
"chosen approach for an interactive educational tool.")

doc = Document(TPL)
# Heading 3 -> Arial
h3s = doc.styles['Heading 3']; h3s.font.name='Arial'; h3s.font.size=Pt(12); h3s.font.bold=True; h3s.font.color.rgb=RGBColor(0,0,0)

def delete_para(p): p._p.getparent().remove(p._p)
def set_para_text(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]: r._r.getparent().remove(r._r)
    else: p.add_run(text)

# ---- title page edits ----
del_anns = []
for p in doc.paragraphs:
    t = p.text.strip()
    if p.style.name=='Heading 1' and t.upper().startswith('INTRODUCERE'): break
    if t.startswith('TITLUL LUCRĂRII'):
        set_para_text(p, TITLE)
        for r in p.runs: r.font.size=Pt(20); r.bold=True
    elif t.startswith('(Arial'): del_anns.append(p)
    elif t.startswith('Candidat:'):
        set_para_text(p, f"Candidat: {AUTHOR}")
        for r in p.runs: r.font.size=Pt(14); r.bold=True
    elif t.startswith('Coordonator'):
        set_para_text(p, f"Coordonator științific: {COORD}")
        for r in p.runs: r.font.size=Pt(14); r.bold=True
    elif t.startswith('Sesiunea:'):
        set_para_text(p, "Sesiunea: Iunie 2026")
        for r in p.runs: r.font.size=Pt(14); r.bold=False
    elif t.startswith('Rezumatul este destinat'):
        set_para_text(p, REZUMAT); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs: r.font.size=Pt(12); r.italic=False
    elif t.startswith('The English abstract'):
        set_para_text(p, ABSTRACT); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
        for r in p.runs: r.font.size=Pt(12); r.italic=False
for p in del_anns: delete_para(p)

# ---- delete instructional body from INTRODUCERE onward ----
dm=False
for p in list(doc.paragraphs):
    if not dm and p.style.name=='Heading 1' and p.text.strip().upper().startswith('INTRODUCERE'): dm=True
    if dm: delete_para(p)

# ---- remove leftover template tables (example table from instructions) ----
for tbl in list(doc.tables):
    tbl._tbl.getparent().remove(tbl._tbl)

# ---- headers (replace at text-node level; placeholders sit in single w:t nodes) ----
hmap = {'Denumire program de studii': f"Universitatea Politehnica Timișoara — {PROGRAM}",
        'Anul universitar': f"Anul universitar {YEAR}", 'Nume candidat': AUTHOR_HDR,
        'Titlul lucrării': TITLE_SHORT}
def fix_header_part(part_elem):
    # work per-paragraph so placeholders split across runs still match
    for wp in part_elem.iter(qn('w:p')):
        ts = wp.findall('.//'+qn('w:t'))
        if not ts: continue
        joined = ''.join(t.text or '' for t in ts)
        new = joined
        for k,v in hmap.items():
            if k in new: new = new.replace(k,v)
        if new != joined:
            ts[0].text = new
            ts[0].set(qn('xml:space'),'preserve')
            for t in ts[1:]: t.text = ''
for sec in doc.sections:
    for hdr in (sec.header, sec.first_page_header, sec.even_page_header):
        if hdr is not None and hdr._element is not None:
            fix_header_part(hdr._element)

# ================= CONTENT HELPERS =================
FIG={"n":0}; TAB={"n":0}; figlist=[]; tablist=[]
def add_cuprins():
    doc.add_page_break()
    p=doc.add_paragraph(style='Heading 1'); r=p.add_run("CUPRINS")
    r.font.name='Arial'; r.font.size=Pt(14); r.bold=True
    tp=doc.add_paragraph(); run=tp.add_run()
    fb=OxmlElement('w:fldChar'); fb.set(qn('w:fldCharType'),'begin')
    it=OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text='TOC \\o "1-3" \\h \\z \\u'
    fs=OxmlElement('w:fldChar'); fs.set(qn('w:fldCharType'),'separate')
    tt=OxmlElement('w:t'); tt.text='Click dreapta aici → „Update Field” pentru a genera cuprinsul.'
    fe=OxmlElement('w:fldChar'); fe.set(qn('w:fldCharType'),'end')
    for el in (fb,it,fs,tt,fe): run._r.append(el)
def body(t, justify=True, indent=True):
    p=doc.add_paragraph(style='Normal')
    p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if indent: p.paragraph_format.first_line_indent=Cm(1.0)
    r=p.add_run(t); r.font.name='Arial'; r.font.size=Pt(12); return p
def bullet(t):
    p=doc.add_paragraph(style='Normal'); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent=Cm(1.0); p.paragraph_format.first_line_indent=Cm(-0.5)
    r=p.add_run("•  "+t); r.font.name='Arial'; r.font.size=Pt(12); return p
def H1(num,txt):
    doc.add_page_break()
    label = f"{num}. {txt.upper()}" if num else txt.upper()
    p=doc.add_paragraph(style='Heading 1'); r=p.add_run(label)
    r.font.name='Arial'; r.font.size=Pt(14); r.bold=True; return p
def H2(num,txt):
    p=doc.add_paragraph(style='Heading 2'); r=p.add_run(f"{num} {txt.upper()}")
    r.font.name='Arial'; r.font.size=Pt(12); r.bold=True; return p
def H3(num,txt):
    p=doc.add_paragraph(style='Heading 3'); r=p.add_run(f"{num} {txt}")
    r.font.name='Arial'; r.font.size=Pt(12); r.bold=True; return p
def fig_img(fname, caption, width=14):
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(FIGS,fname), width=Cm(width))
    FIG["n"]+=1; n=FIG["n"]
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.add_run(f"Figura {n} – {caption}"); r.font.name='Arial'; r.font.size=Pt(10); r.italic=True
    figlist.append((n,caption)); doc.add_paragraph(); return n
def fig_ph(caption):
    doc.add_paragraph()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f"[ Captură de ecran – {caption} ]"); r.font.name='Arial'; r.font.size=Pt(11)
    r.italic=True; r.font.color.rgb=RGBColor(0x88,0x88,0x88)
    FIG["n"]+=1; n=FIG["n"]
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rr=c.add_run(f"Figura {n} – {caption}"); rr.font.name='Arial'; rr.font.size=Pt(10); rr.italic=True
    figlist.append((n,caption)); doc.add_paragraph(); return n
def code(codetxt, caption):
    for line in codetxt.split('\n'):
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.6)
        p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
        r=p.add_run(line if line else ' '); r.font.name='Consolas'; r.font.size=Pt(9)
    FIG["n"]+=1; n=FIG["n"]
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.add_run(f"Figura {n} – {caption}"); r.font.name='Arial'; r.font.size=Pt(10); r.italic=True
    figlist.append((n,caption)); doc.add_paragraph(); return n
def table_cap(caption):
    TAB["n"]+=1; n=TAB["n"]
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.add_run(f"Tabelul {n}. {caption}"); r.font.name='Arial'; r.font.size=Pt(10); r.italic=True
    tablist.append((n,caption)); return n

add_cuprins()
import content
content.build(doc, dict(body=body,bullet=bullet,H1=H1,H2=H2,H3=H3,fig_img=fig_img,
                        fig_ph=fig_ph,code=code,table_cap=table_cap,doc=doc,
                        set_cell=None,figlist=figlist,tablist=tablist,AUTHOR=AUTHOR))

doc.save(LOCAL)
print("SAVED LOCAL:", LOCAL)
saved_to = OUT
try:
    doc.save(OUT)
    print("SAVED:", OUT)
except PermissionError:
    import time
    alt = OUT.replace('.docx', f'_{time.strftime("%H%M")}.docx')
    doc.save(alt); saved_to = alt
    print("Destinația era blocată; salvat ca:", alt)
from docx import Document as D
d=D(LOCAL); words=sum(len(p.text.split()) for p in d.paragraphs)
print(f"paras={len(d.paragraphs)} tables={len(d.tables)} ~words={words} figs={FIG['n']}")
