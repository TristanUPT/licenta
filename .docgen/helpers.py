# -*- coding: utf-8 -*-
"""Shared formatting helpers for the ResoLab thesis generator."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PURPLE = RGBColor(0x7c, 0x3a, 0xed)
DARKGREY = RGBColor(0x22, 0x22, 0x22)
CODEBG = RGBColor(0xF3, 0xF3, 0xF3)

FIG = {"n": 0}
TAB = {"n": 0}
fig_index = []   # (num, caption, )
tab_index = []

def new_doc():
    doc = Document()
    # Base style
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_after = Pt(6)
    # Margins
    for s in doc.sections:
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    return doc

def set_cell_text(cell, text, bold=False, size=11, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER}[align]
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold
    return cell

def body(doc, text, justify=True, first_indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    return p

def h1(doc, num, text):
    doc.add_page_break()
    p = doc.add_heading(level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{num} {text.upper()}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.bold = True
    run.font.color.rgb = DARKGREY
    return p

def h2(doc, num, text):
    p = doc.add_heading(level=2)
    run = p.add_run(f"{num} {text.upper()}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = True
    run.font.color.rgb = DARKGREY
    return p

def h3(doc, num, text):
    p = doc.add_heading(level=3)
    run = p.add_run(f"{num} {text}")
    run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = True
    run.font.color.rgb = DARKGREY
    return p

def figure(doc, caption):
    FIG["n"] += 1
    n = FIG["n"]
    ph = doc.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ph.add_run(f"[FIGURA {n} – {caption}]")
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(f"Fig. {n} – {caption}")
    rc.font.name = 'Times New Roman'; rc.font.size = Pt(10); rc.italic = True
    fig_index.append((n, caption))
    return n

def code_block(doc, code, caption):
    FIG["n"] += 1
    n = FIG["n"]
    for line in code.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line if line else ' ')
        r.font.name = 'Consolas'; r.font.size = Pt(9)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(f"Fig. {n} – {caption}")
    rc.font.name = 'Times New Roman'; rc.font.size = Pt(10); rc.italic = True
    fig_index.append((n, caption))
    return n

def table_caption(doc, caption):
    TAB["n"] += 1
    n = TAB["n"]
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = cap.add_run(f"Tabel {n} – {caption}")
    rc.font.name = 'Times New Roman'; rc.font.size = Pt(10); rc.italic = True
    tab_index.append((n, caption))
    return n

def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def running_header(doc, author, title):
    sec = doc.sections[0]
    hdr = sec.header
    hp = hdr.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(f"Informatică  ·  2024–2025\n{author} – {title}")
    r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0x77,0x77,0x77)
    ftr = sec.footer
    fp = ftr.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pre = fp.add_run("- "); pre.font.size = Pt(10); pre.font.name='Times New Roman'
    add_page_number_field(fp)
    post = fp.add_run(" -"); post.font.size = Pt(10); post.font.name='Times New Roman'

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = "Actualizați câmpul (click dreapta → Update Field) pentru a genera cuprinsul."
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar); run._r.append(instrText); run._r.append(fldChar2)
    run._r.append(t); run._r.append(fldChar3)
