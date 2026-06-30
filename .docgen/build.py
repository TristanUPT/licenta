# -*- coding: utf-8 -*-
import sys
sys.path.insert(0, '.')
from helpers import new_doc, running_header
import part1, part2, part3

AUTHOR = "Tristan Cvasniuc"
TITLE = "ResoLab – mini-DAW educațional"

doc = new_doc()
running_header(doc, AUTHOR, TITLE)

# Front matter
part1.front_matter(doc)
part1.chapter1(doc)
part1.chapter2(doc)
part1.chapter3(doc)
# Core
part2.chapter4(doc)
part2.chapter5(doc)
# Back
part3.chapter6(doc)
part3.chapter7(doc)
part3.bibliography(doc)
part3.back_matter(doc)

out = "/mnt/c/Users/Tup20/OneDrive - Universitatea Politehnica Timisoara/Desktop/documentatii licenta/Documentatie_Licenta_ResoLab.docx"
doc.save(out)
# Also save a local copy
doc.save("/home/tup20/licenta/Documentatie_Licenta_ResoLab.docx")
print("SAVED to:")
print(out)

# quick stats
from docx import Document
d = Document("/home/tup20/licenta/Documentatie_Licenta_ResoLab.docx")
words = sum(len(p.text.split()) for p in d.paragraphs)
print(f"Paragraphs: {len(d.paragraphs)}  Tables: {len(d.tables)}  ~Words: {words}")
