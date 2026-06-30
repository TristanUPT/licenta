# -*- coding: utf-8 -*-
from helpers import *
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

AUTHOR = "Tristan CVASNIUC"
TITLE = "ResoLab – mini-DAW educațional pentru procesarea audio în timp real în browser"
COORD = "Ș.l.dr.inf. Cristian ZIMBRU"

def front_matter(doc):
    # ---- Title page ----
    def tp(text, size, bold=False, color=None, after=6, before=0, align='center'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
        r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold
        if color: r.font.color.rgb = color
        return p
    tp("UNIVERSITATEA POLITEHNICA TIMIȘOARA", 13, True)
    tp("FACULTATEA DE AUTOMATICĂ ȘI CALCULATOARE", 12, True)
    tp("Specializarea Informatică", 12)
    tp("2024–2025", 12, after=120)
    tp("LUCRARE DE LICENȚĂ", 16, True, after=60)
    tp("ResoLab – MINI-DAW EDUCAȚIONAL PENTRU", 16, True, PURPLE, after=0)
    tp("PROCESAREA AUDIO ÎN TIMP REAL ÎN BROWSER", 16, True, PURPLE, after=120)
    tp("FOLOSIND RUST ȘI WEBASSEMBLY", 14, True, after=160)
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(2.5)
    for lbl, val in [("Candidat: ", AUTHOR), ("Coordonator științific: ", COORD)]:
        pp = doc.add_paragraph(); pp.alignment=WD_ALIGN_PARAGRAPH.LEFT
        pp.paragraph_format.left_indent = Cm(2.5)
        r1 = pp.add_run(lbl); r1.bold=True; r1.font.size=Pt(12); r1.font.name='Times New Roman'
        r2 = pp.add_run(val); r2.font.size=Pt(12); r2.font.name='Times New Roman'
    tp("Sesiunea: Iulie 2025", 12, before=160)
    tp("Timișoara", 12, True, before=120)

    # ---- Rezumat ----
    doc.add_page_break()
    hp = doc.add_heading(level=1); 
    r = hp.add_run("REZUMAT"); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True; r.font.color.rgb=DARKGREY
    body(doc, "Lucrarea de față descrie ResoLab, o aplicație web destinată învățării și exersării prelucrării sunetului, care rulează integral în browser, fără instalare și fără server de procesare. Spre deosebire de editoarele audio clasice, în care utilizatorul aplică efecte fără a înțelege ce se întâmplă cu semnalul, ResoLab a fost gândit pentru a face vizibil și explicabil fiecare pas al procesării. Fiecare parametru este însoțit de explicații în două registre — pentru începători și pentru utilizatori avansați — iar răspunsul sonor este dublat de reprezentări grafice actualizate în timp real.")
    body(doc, "Nucleul de procesare a semnalului este scris în limbajul Rust și compilat către WebAssembly, ceea ce permite atingerea unei latențe reduse direct în pagina web. Acest motor rulează izolat, pe firul de execuție audio al browserului, printr-un AudioWorklet care funcționează ca punte subțire între interfața grafică și codul nativ. Interfața este construită cu React și TypeScript, iar starea aplicației este coordonată prin biblioteca Zustand, organizată pe domenii independente.")
    body(doc, "Aplicația oferă un lanț de efecte reconfigurabil format din șaisprezece procesoare audio (compresor, egalizator parametric, reverberație, întârziere, poartă de zgomot, limitator și altele), un sintetizator substractiv polifonic cu opt voci, un sistem de înregistrare cu selecție a interfeței audio externe, un metronom cu programare temporală precisă, vizualizări spectrale și un strat educațional care analizează configurația curentă și propune observații și recomandări. Rezultatul prelucrării poate fi exportat ca fișier WAV. În urma testelor efectuate, motorul a procesat lanțuri complexe de efecte fără întreruperi audibile, confirmând că abordarea bazată pe WebAssembly este viabilă pentru un instrument educațional interactiv.")

    # ---- Abstract ----
    doc.add_page_break()
    hp = doc.add_heading(level=1)
    r = hp.add_run("ABSTRACT"); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True; r.font.color.rgb=DARKGREY
    body(doc, "This thesis presents ResoLab, a web application for learning and practising sound engineering that runs entirely in the browser, with no installation and no processing server. Unlike conventional audio editors, where the user applies effects without grasping what happens to the signal, ResoLab was designed to make every processing step visible and explainable. Each parameter is accompanied by explanations at two levels — beginner and advanced — while the sonic result is mirrored by graphical representations updated in real time.")
    body(doc, "The signal-processing core is written in Rust and compiled to WebAssembly, which makes low-latency processing possible directly inside the web page. This engine runs in isolation, on the browser audio thread, through an AudioWorklet that acts as a thin bridge between the graphical interface and the native code. The interface is built with React and TypeScript, and application state is coordinated through the Zustand library, organised into independent domains.")
    body(doc, "The application provides a reconfigurable effects chain of sixteen audio processors (compressor, parametric equaliser, reverb, delay, noise gate, limiter and others), an eight-voice polyphonic subtractive synthesiser, a recording system with external audio-interface selection, a metronome with precise temporal scheduling, spectral visualisations and an educational layer that analyses the current configuration and offers observations and recommendations. The processed result can be exported as a WAV file. Testing confirmed that the engine processed complex effect chains without audible dropouts, validating the WebAssembly approach for an interactive educational tool.")

    # ---- Cuprins ----
    doc.add_page_break()
    hp = doc.add_heading(level=1)
    r = hp.add_run("CUPRINS"); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True; r.font.color.rgb=DARKGREY
    add_toc(doc)

def chapter1(doc):
    h1(doc, "1", "Introducere")
    body(doc, "Acest capitol introductiv așază tema lucrării în contextul ei mai larg, explică motivele care au condus la alegerea ei și enunță obiectivele urmărite pe parcursul dezvoltării. În finalul capitolului este descrisă pe scurt organizarea întregii lucrări, astfel încât cititorul să poată parcurge ulterior fiecare secțiune cu o imagine clară asupra rolului pe care îl ocupă.")

    h2(doc, "1.1", "Domeniul proiectului")
    body(doc, "Prelucrarea digitală a sunetului a încetat de mult să fie apanajul studiourilor profesionale. Telefonul din buzunar înregistrează, montează și publică material audio, iar platformele de distribuție au coborât bariera de intrare aproape la zero. În acest peisaj, decalajul nu mai este unul de acces la unelte, ci unul de înțelegere: foarte mulți utilizatori manevrează un compresor sau un egalizator fără să poată explica ce transformă acestea în semnal. Lucrarea se înscrie în zona aplicațiilor educaționale interactive pentru ingineria sunetului, la intersecția dintre prelucrarea numerică a semnalelor și tehnologiile web moderne.")
    body(doc, "Un program de tip stație de lucru audio digitală — cunoscut în literatura de specialitate prin acronimul DAW, de la „digital audio workstation” — reunește înregistrarea, editarea, procesarea și mixarea sunetului într-un mediu unitar. Asemenea programe sunt, în mod tradițional, aplicații native voluminoase, care presupun instalare, licențiere și un calculator suficient de puternic. Apariția standardului WebAssembly a schimbat această ecuație, întrucât a făcut posibilă rularea în browser a unui cod compilat, apropiat ca viteză de execuția nativă [1]. ResoLab valorifică tocmai această deschidere, mutând procesarea audio de performanță în pagina web.")
    body(doc, "Proiectul nu își propune să concureze instrumentele profesionale pe terenul numărului de funcții, ci să acopere o nișă distinctă: aceea a învățării asistate. Accentul cade pe transparența procesării — utilizatorul vede, aude și citește, în același timp, ce face fiecare efect. Din acest motiv, fiecare decizie de proiectare a fost subordonată scopului didactic, de la limbajul dual al explicațiilor până la vizualizările care însoțesc redarea.")

    h2(doc, "1.2", "Specificațiile generale ale aplicației")
    body(doc, "ResoLab funcționează ca un atelier audio complet, accesibil printr-un simplu acces la o adresă web. Utilizatorul încarcă un fișier sonor prin tragere și plasare sau alege unul dintre eșantioanele demonstrative incluse, după care construiește un lanț de efecte pe care îl poate reordona, dezactiva selectiv sau compara instantaneu cu semnalul original. Întreaga procesare are loc local, pe dispozitivul utilizatorului; niciun fișier nu este transmis către un server, ceea ce asigură atât confidențialitate, cât și independență față de o conexiune la rețea.")
    body(doc, "Pe lângă prelucrarea fișierelor existente, aplicația include un sintetizator care permite generarea de sunete de la zero, un modul de înregistrare capabil să capteze semnalul de la o interfață audio externă și un metronom util pentru exersarea ritmică. Stratul educațional traversează toate aceste componente: el comentează valorile parametrilor, semnalează configurațiile riscante și sugerează efecte potrivite în funcție de conținutul spectral al semnalului. Limba interfeței poate fi comutată între română și engleză, iar nivelul explicațiilor poate fi adaptat experienței utilizatorului.")

    h2(doc, "1.3", "Obiective")
    body(doc, "Pentru a delimita cu precizie ținta lucrării, am formulat de la început un set de obiective concrete, care au ghidat ulterior toate deciziile tehnice:")
    bullet(doc, "Construirea unui motor de procesare audio în timp real care să ruleze în browser cu latență redusă, scris în Rust și compilat către WebAssembly, capabil să susțină un lanț reconfigurabil de efecte fără întreruperi audibile.")
    bullet(doc, "Proiectarea unei interfețe care nu doar aplică efecte, ci explică funcționarea lor, prin tooltipuri contextuale, feedback dinamic asupra parametrilor și vizualizări sincronizate cu redarea, adresate deopotrivă începătorilor și utilizatorilor avansați.")
    bullet(doc, "Implementarea unui set reprezentativ de procesoare audio — compresor, egalizator parametric, reverberație, întârziere și altele — pornind de la algoritmi consacrați în literatura de specialitate, astfel încât comportamentul lor să fie corect din punct de vedere al ingineriei sunetului.")
    bullet(doc, "Asimilarea, în regim autodidact, a unui ansamblu de tehnologii noi pentru autor — limbajul Rust, modelul de execuție WebAssembly și interfața de programare Web Audio — și integrarea lor coerentă într-un produs funcțional.")
    bullet(doc, "Asigurarea unei experiențe complete de la capăt la capăt: încărcarea sunetului, procesarea lui, vizualizarea rezultatului și exportul într-un fișier WAV valid, redabil în orice player extern.")

    h2(doc, "1.4", "Structura lucrării")
    body(doc, "Lucrarea este organizată în șapte capitole. Capitolul de față introduce tema, motivația și obiectivele. Al doilea capitol analizează stadiul actual al domeniului, prin examinarea unor aplicații asemănătoare și prin compararea lor sistematică cu soluția propusă. Al treilea capitol prezintă fundamentele teoretice și tehnologiile pe care se sprijină dezvoltarea. Al patrulea capitol descrie soluția propusă la nivel arhitectural, împreună cu principalele cazuri de utilizare. Al cincilea capitol intră în detaliile de implementare ale fiecărei funcționalități, însoțite de fragmente de cod relevante. Al șaselea capitol prezintă modul de utilizare a aplicației din perspectiva utilizatorului final, iar ultimul capitol formulează concluziile și conturează direcțiile de dezvoltare ulterioară.")

def chapter2(doc):
    h1(doc, "2", "Analiza stadiului actual în domeniu")
    body(doc, "Înainte de a justifica oportunitatea unei noi aplicații, este firesc să examinăm ce oferă deja peisajul instrumentelor existente. În acest capitol sunt analizate patru soluții reprezentative pentru editarea și producția audio în browser sau pe desktop, urmărind cu precădere funcționalitățile pe care le pun la dispoziție și tehnologiile pe care se sprijină. Capitolul se încheie cu o comparație sistematică, sintetizată într-un tabel, din care reiese spațiul pe care ResoLab îl ocupă.")

    h2(doc, "2.1", "Audacity și portarea Wavacity")
    body(doc, "Audacity este probabil cel mai cunoscut editor audio gratuit și cu sursă deschisă, folosit de peste două decenii pentru înregistrare și editare multi-pistă [2]. Programul oferă un arsenal bogat de efecte, suport pentru numeroase formate și o comunitate vastă. Limitarea sa, din perspectiva temei de față, este de natură pedagogică: interfața expune efectele sub formă de ferestre de dialog cu parametri numerici, fără a explica utilizatorului ce transformă fiecare reglaj asupra semnalului. Audacity este o unealtă de producție, nu un mediu de învățare.")
    body(doc, "Proiectul Wavacity reprezintă o portare a Audacity în browser, realizată prin compilarea codului C++ original către WebAssembly [3]. Această inițiativă demonstrează, în mod concret, că un editor audio complex poate rula într-o pagină web, validând astfel premisa tehnologică pe care se sprijină și ResoLab. Totuși, Wavacity moștenește integral modelul de interacțiune al aplicației originale, inclusiv absența unui strat explicativ. [FIGURA placeholder este descrisă mai jos.]")
    figure(doc, "Interfața de editare a aplicației Audacity, cu fereastra de dialog a unui efect")

    h2(doc, "2.2", "Soundtrap")
    body(doc, "Soundtrap, parte a ecosistemului Spotify, este o stație de lucru audio colaborativă care funcționează integral în browser [4]. Aplicația permite înregistrarea, aranjarea pe piste și aplicarea de efecte, punând un accent deosebit pe colaborarea în timp real între mai mulți utilizatori și pe integrarea cu serviciile educaționale. Soundtrap include resurse didactice și se adresează inclusiv mediului școlar, fapt care o apropie tematic de proiectul de față.")
    body(doc, "Diferența esențială ține de natura procesării. Soundtrap se bazează pe infrastructura cloud și pe un model de abonament, iar efectele sunt prezentate ca unelte gata făcute, fără a deschide o fereastră asupra mecanismelor interne. Componenta educațională vizează compoziția și producția muzicală, nu înțelegerea aprofundată a procesării de semnal la nivel de parametru.")

    h2(doc, "2.3", "BandLab")
    body(doc, "BandLab este o platformă gratuită de creație muzicală, disponibilă în browser și pe dispozitive mobile, care combină o stație de lucru audio cu o rețea socială dedicată muzicienilor [5]. Utilizatorii pot înregistra, mixa și publica piese, dispunând de o bibliotecă generoasă de instrumente virtuale și bucle. Modelul gratuit și accesibilitatea au făcut din BandLab o opțiune populară printre creatorii începători.")
    body(doc, "Și în acest caz, orientarea este către rezultatul muzical, nu către procesul de învățare a ingineriei sunetului. Efectele sunt aplicate prin interfețe simplificate, optimizate pentru rapiditate, iar utilizatorul nu primește explicații despre funcționarea lor. Procesarea se sprijină în bună măsură pe servere, ceea ce presupune o conexiune permanentă.")

    h2(doc, "2.4", "AudioMass")
    body(doc, "AudioMass este un editor audio gratuit, cu sursă deschisă, care rulează în browser fără a necesita instalare sau cont [6]. Spre deosebire de soluțiile anterioare, AudioMass este construit integral pe interfața de programare Web Audio nativă a browserului, fără un motor compilat separat. Aplicația oferă funcții de tăiere, aplicare de efecte și vizualizare a formei de undă, într-o interfață ușoară și rapidă.")
    body(doc, "AudioMass ilustrează abordarea opusă celei alese în ResoLab: procesarea se face prin nodurile native ale interfeței Web Audio. Această cale este simplă, însă oferă un control limitat asupra algoritmilor și nu permite implementarea propriilor procesoare cu precizia dorită. Lipsa unui strat educațional o menține, de asemenea, în categoria uneltelor de editare.")

    h2(doc, "2.5", "Comparație și concluzii")
    body(doc, "Punând alături soluțiile analizate, se conturează un tipar clar: instrumentele existente sunt fie editoare orientate către producție, fie platforme de creație muzicală, niciuna nefiind proiectată în jurul învățării procesării de semnal. Tabelul următor sintetizează prezența sau absența unor caracteristici relevante pentru tema lucrării.")
    table_caption(doc, "Comparație între aplicațiile analizate și ResoLab")
    rows = [
        ["Caracteristică", "Audacity / Wavacity", "Soundtrap", "BandLab", "AudioMass", "ResoLab"],
        ["Rulează în browser", "Parțial / Da", "Da", "Da", "Da", "Da"],
        ["Procesare locală (fără server)", "Da", "Nu", "Nu", "Da", "Da"],
        ["Motor DSP propriu (Rust/WASM)", "Da (C++)", "Nu", "Nu", "Nu", "Da"],
        ["Strat educațional / explicații", "Nu", "Parțial", "Nu", "Nu", "Da"],
        ["Feedback dinamic pe parametri", "Nu", "Nu", "Nu", "Nu", "Da"],
        ["Recomandări pe analiză spectrală", "Nu", "Nu", "Nu", "Nu", "Da"],
        ["Vizualizări în timp real", "Parțial", "Parțial", "Parțial", "Da", "Da"],
        ["Sintetizator integrat", "Nu", "Da", "Da", "Nu", "Da"],
        ["Fără cont / fără abonament", "Da", "Nu", "Parțial", "Da", "Da"],
    ]
    t = doc.add_table(rows=len(rows), cols=6)
    t.style = 'Table Grid'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            set_cell_text(t.rows[i].cells[j], val, bold=(i==0 or j==5), size=9,
                          align='left' if j==0 else 'center')
    body(doc, "Din tabel reiese că niciuna dintre aplicațiile consacrate nu îmbină procesarea locală de performanță cu un strat explicit de învățare. Soluțiile care rulează local nu explică procesarea, iar cele care includ resurse educaționale o fac la nivelul compoziției și depind de infrastructura cloud. ResoLab ocupă tocmai acest spațiu liber: un motor de procesare propriu, rulat local, dublat de un strat educațional care comentează și recomandă în funcție de starea reală a semnalului. Această combinație constituie principala contribuție a lucrării față de stadiul actual al domeniului.")

def chapter3(doc):
    h1(doc, "3", "Fundamente teoretice")
    body(doc, "Acest capitol trece în revistă tehnologiile și bibliotecile pe care se sprijină dezvoltarea aplicației, oferind contextul necesar pentru înțelegerea deciziilor descrise în capitolele următoare. Prezentarea se concentrează asupra elementelor cu adevărat relevante pentru proiect, fără a epuiza documentația fiecărui instrument în parte.")

    h2(doc, "3.1", "Rust și WebAssembly")
    body(doc, "Rust este un limbaj de programare de sistem orientat către siguranță și performanță, care garantează absența unei categorii întregi de erori de memorie fără a recurge la un colector de gunoaie [7]. Tocmai această combinație îl face potrivit pentru cod audio: procesarea de semnal cere viteză predictibilă, iar pauzele introduse de un colector de gunoaie ar provoca întreruperi audibile. În ResoLab, întregul nucleu de prelucrare a semnalului este scris în Rust.")
    body(doc, "WebAssembly este un format binar de instrucțiuni de nivel scăzut, conceput ca țintă de compilare portabilă pentru limbaje precum C, C++ și Rust, executabil în browser la o viteză apropiată de cea nativă [1]. Codul Rust al motorului este compilat către acest format cu ajutorul utilitarului wasm-pack, rezultând un modul care poate fi încărcat și apelat din JavaScript. Comunicarea la granița dintre cele două lumi se realizează prin funcții exportate cu convenția C și prin schimb de pointeri către memoria liniară a modulului, evitând copierea inutilă a tampoanelor de eșantioane.")

    h2(doc, "3.2", "Interfața de programare Web Audio și AudioWorklet")
    body(doc, "Interfața Web Audio oferă browserului un graf de procesare a sunetului, alcătuit din noduri interconectate, prin care semnalul curge de la sursă către ieșirea audio [8]. Pentru procesarea personalizată, standardul pune la dispoziție mecanismul AudioWorklet, care permite execuția de cod propriu pe firul dedicat audio, separat de firul principal al interfeței. Acest fir rulează cu prioritate ridicată, ceea ce reduce riscul întreruperilor.")
    body(doc, "În arhitectura aplicației, procesorul AudioWorklet îndeplinește rolul unei punți subțiri: el încarcă modulul WebAssembly, copiază tampoanele de intrare și ieșire și transmite mesajele de control către motorul nativ. Decizia de a păstra acest strat cât mai subțire este deliberată — orice logică suplimentară pe firul audio ar putea introduce întârzieri. Procesarea propriu-zisă se desfășoară în blocuri de o sută douăzeci și opt de eșantioane, dimensiunea standard a cuantei de redare din Web Audio.")

    h2(doc, "3.3", "React și TypeScript")
    body(doc, "Interfața grafică este construită cu React, o bibliotecă pentru construirea de interfețe pe bază de componente, în care reprezentarea vizuală este o funcție a stării [9]. Modelul declarativ al React se potrivește bine unei aplicații cu multe controale interdependente, deoarece sincronizarea dintre starea internă și ceea ce vede utilizatorul este gestionată automat. Toate componentele aplicației sunt funcționale și folosesc mecanismul de hookuri.")
    body(doc, "Peste React am folosit TypeScript, o extensie a limbajului JavaScript care adaugă un sistem de tipuri verificat la compilare [10]. Tipizarea statică s-a dovedit deosebit de utilă într-un proiect cu numeroase mesaje schimbate între interfață și motorul audio, deoarece a permis surprinderea timpurie a nepotrivirilor de structură. Codul a fost scris cu opțiunile stricte de verificare activate, ceea ce a redus considerabil clasa erorilor strecurate la execuție.")

    h2(doc, "3.4", "Gestiunea stării cu Zustand")
    body(doc, "Pentru coordonarea stării globale am ales biblioteca Zustand, o soluție minimalistă care expune starea prin depozite independente, accesate cu ajutorul unor selectoare [11]. Spre deosebire de abordările bazate pe un depozit unic și monolitic, Zustand încurajează împărțirea stării pe domenii. În ResoLab există depozite distincte pentru lanțul de efecte, pentru fișierul audio și redare, pentru datele de analiză, pentru starea educațională, pentru preseturi, pentru interfață și pentru sintetizator.")
    body(doc, "Această separare aduce două avantaje practice. Pe de o parte, fiecare componentă se abonează doar la felia de stare de care are nevoie, ceea ce limitează redesenările inutile. Pe de altă parte, logica de modificare a stării este definită în interiorul depozitelor, sub formă de acțiuni, păstrând componentele curate și ușor de urmărit.")

    h2(doc, "3.5", "Biblioteci pentru interfață și vizualizare")
    body(doc, "Aspectul vizual se sprijină pe TailwindCSS, un sistem de stilizare bazat pe clase utilitare, completat de primitivele accesibile oferite de Radix UI pentru elemente precum ferestre de dialog, meniuri și comutatoare [12]. Reprezentarea formei de undă este realizată cu biblioteca wavesurfer.js, care desenează semnalul și gestionează interacțiunea de selecție a regiunilor de buclă [13]. Pentru afișarea spectrului în timp real am folosit audiomotion-analyzer, iar pentru programarea temporală a fost integrată biblioteca Tone.js, utilizată exclusiv pentru funcții de transport, nu pentru procesarea efectelor. Persistența preseturilor se realizează printr-o bază de date IndexedDB, accesată prin biblioteca idb, în timp ce preferințele de interfață sunt salvate în spațiul local de stocare al browserului.")
