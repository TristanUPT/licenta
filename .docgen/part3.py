# -*- coding: utf-8 -*-
from helpers import *
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def chapter6(doc):
    h1(doc, "6", "Utilizarea aplicației")
    body(doc, "Acest capitol prezintă aplicația din perspectiva utilizatorului final, urmărind pașii pe care îi parcurge de la primul contact cu interfața până la obținerea unui rezultat exportat. Descrierea însoțește principalele ecrane și acțiuni, oferind o imagine de ansamblu asupra modului de operare.")

    h2(doc, "6.1", "Încărcarea unui fișier și redarea")
    body(doc, "La deschiderea aplicației, utilizatorul întâmpină un spațiu central pregătit să primească un fișier audio. Acesta poate fi adus prin tragere și plasare sau selectat cu un clic, iar pentru explorare rapidă sunt disponibile câteva eșantioane demonstrative în panoul din stânga. Odată încărcat, fișierul este afișat sub forma undei sonore, iar bara de transport din partea de jos devine activă, permițând redarea, oprirea și parcurgerea materialului.")
    figure(doc, "Ecranul inițial cu zona de încărcare a fișierului audio")
    body(doc, "Pe forma de undă, utilizatorul poate fixa o regiune de buclă prin tragere, pentru a relua un anumit fragment, și poate sări într-un punct anume printr-un simplu clic. Aceste interacțiuni sunt dublate de scurtături de tastatură, care accelerează lucrul pentru utilizatorii obișnuiți.")

    h2(doc, "6.2", "Construirea lanțului de efecte")
    body(doc, "Adăugarea unui efect se face din butonul dedicat, care deschide lista procesoarelor disponibile. Efectele se așază în ordine, formând un lanț pe care utilizatorul îl poate rearanja prin tragere, folosind mânerul afișat pe fiecare card. Fiecare efect poate fi dezactivat individual, iar întregul lanț poate fi ocolit dintr-o singură comandă, pentru a compara rapid semnalul procesat cu cel original.")
    figure(doc, "Lanțul de efecte cu mai multe procesoare active")
    body(doc, "Reglarea parametrilor se face prin butoane rotative și culisoare. La trecerea cursorului peste oricare dintre ele, apare o explicație contextuală, iar panoul din dreapta comentează configurația și propune ajustări. Orice modificare poate fi anulată sau refăcută, ceea ce încurajează experimentarea fără teama de a strica rezultatul.")

    h2(doc, "6.3", "Salvarea și încărcarea preseturilor")
    body(doc, "Configurațiile reușite pot fi salvate ca preseturi personalizate, cu un nume ales de utilizator. Acestea sunt păstrate local, într-o bază de date a browserului, astfel încât rămân disponibile și la vizitele ulterioare. Pe lângă preseturile proprii, aplicația include un set de configurații de fabrică, gândite ca puncte de plecare pentru situații frecvente.")

    h2(doc, "6.4", "Exportul rezultatului")
    body(doc, "Când rezultatul este satisfăcător, utilizatorul îl poate exporta într-un fișier WAV printr-o singură comandă din bara de transport. Procesarea se desfășoară în fundal, iar fișierul descărcat poate fi redat în orice player extern, cu durata și calitatea corecte. Astfel se închide ciclul complet de lucru, de la sunetul brut la rezultatul finit.")
    figure(doc, "Exportul rezultatului procesat într-un fișier WAV")

def chapter7(doc):
    h1(doc, "7", "Concluzii")
    body(doc, "În urma parcurgerii întregului proces de proiectare și dezvoltare, capitolul de față sintetizează rezultatele obținute, problemele întâmpinate și modul în care au fost depășite, plusurile aplicației față de soluțiile existente și direcțiile pe care le considerăm promițătoare pentru viitor.")

    h2(doc, "7.1", "Rezultate obținute")
    body(doc, "Obiectivele formulate la începutul lucrării au fost atinse. Am construit un motor de procesare audio care rulează în browser cu latență redusă, scris în Rust și compilat către WebAssembly, capabil să susțină un lanț reconfigurabil de șaisprezece efecte fără întreruperi audibile. Am realizat o interfață care nu doar aplică efecte, ci le explică, prin tooltipuri, feedback dinamic și vizualizări sincronizate, adaptate atât începătorilor, cât și utilizatorilor avansați. Procesoarele audio au fost implementate pornind de la algoritmi consacrați, ceea ce le conferă un comportament corect din punct de vedere al ingineriei sunetului. Nu în ultimul rând, am asigurat un ciclu complet de lucru, de la încărcarea sunetului până la exportul într-un fișier WAV valid.")
    body(doc, "Dincolo de produsul în sine, procesul a presupus asimilarea, în regim autodidact, a unui set consistent de tehnologii noi pentru mine — limbajul Rust, modelul de execuție WebAssembly și interfața Web Audio — și integrarea lor într-un întreg coerent. Această experiență reprezintă, în sine, unul dintre câștigurile cele mai durabile ale lucrării.")

    h2(doc, "7.2", "Probleme întâmpinate și soluțiile adoptate")
    body(doc, "Pe parcursul dezvoltării au apărut mai multe dificultăți tehnice. Comunicarea la granița dintre JavaScript și WebAssembly a impus o atenție deosebită gestiunii memoriei, întrucât creșterea memoriei modulului invalidează vederile către tampoane; soluția a fost reobținerea acestor vederi după fiecare operație care putea provoca o realocare. Precizia temporală a metronomului, afectată inițial de fluctuațiile cronometrelor obișnuite, a fost rezolvată prin tehnica programării anticipate. Durata greșită a fișierelor exportate, cauzată de un antet scris înainte de cunoașterea dimensiunii finale, a fost corectată prin rescrierea antetului cu valorile exacte. Fiecare dintre aceste obstacole a constituit, retrospectiv, o ocazie de aprofundare.")

    h2(doc, "7.3", "Contribuții față de soluțiile existente")
    body(doc, "Comparativ cu aplicațiile analizate în al doilea capitol, ResoLab aduce o combinație pe care niciuna dintre ele nu o oferă: procesare locală de performanță, realizată cu un motor propriu, dublată de un strat educațional care comentează și recomandă în funcție de starea reală a semnalului. Acolo unde editoarele consacrate aplică efecte fără a le explica, iar platformele de creație se sprijină pe infrastructura cloud și vizează compoziția, aplicația de față pune înțelegerea procesării în centrul experienței. Această orientare didactică, susținută tehnic de un motor rapid și transparent, reprezintă contribuția principală a lucrării.")

    h2(doc, "7.4", "Direcții de dezvoltare ulterioară")
    body(doc, "Aplicația rămâne deschisă mai multor direcții de extindere. Programarea temporală a metronomului și a arpegiatorului ar putea fi mutată pe o bază mai precisă, sincronizată cu ceasul audio. Suportul MIDI ar putea fi îmbogățit cu mesaje de control continuu și cu sensibilitate la viteză. Stratul educațional s-ar putea extinde cu lecții interactive structurate și cu exerciții ghidate. De asemenea, exportul fișierelor de mari dimensiuni ar beneficia de un indicator de progres și de o procesare eșalonată. Aceste direcții depășesc scopul lucrării de față, dar conturează un parcurs firesc de maturizare a proiectului.")

def bibliography(doc):
    h1(doc, "8", "Bibliografie")
    refs = [
        ("WebAssembly – Specification and Overview", "https://webassembly.org/", "12.06.2025"),
        ("Audacity – Free, open source audio editor", "https://www.audacityteam.org/", "12.06.2025"),
        ("Wavacity – Audacity ported to the browser via WebAssembly", "https://wavacity.com/", "12.06.2025"),
        ("Soundtrap by Spotify – Online music studio", "https://www.soundtrap.com/", "12.06.2025"),
        ("BandLab – Free music making platform", "https://www.bandlab.com/", "12.06.2025"),
        ("AudioMass – Free, open source web audio editor", "https://audiomass.co/", "12.06.2025"),
        ("The Rust Programming Language – Official Book, Steve Klabnik, Carol Nichols", "https://doc.rust-lang.org/book/", "12.06.2025"),
        ("MDN Web Docs – Web Audio API", "https://developer.mozilla.org/en-US/docs/Web/API/Web_Audio_API", "13.06.2025"),
        ("React – Official Documentation", "https://react.dev/", "13.06.2025"),
        ("TypeScript – Official Documentation", "https://www.typescriptlang.org/docs/", "13.06.2025"),
        ("Zustand – State management for React", "https://zustand.docs.pmnd.rs/", "13.06.2025"),
        ("Tailwind CSS – Documentation; Radix UI Primitives", "https://tailwindcss.com/docs", "13.06.2025"),
        ("wavesurfer.js – Audio waveform visualisation library", "https://wavesurfer.xyz/", "13.06.2025"),
        ("Udo Zölzer (ed.), „DAFX: Digital Audio Effects”, 2nd Edition, Wiley, 2011", "https://www.dafx.de/", "14.06.2025"),
        ("Robert Bristow-Johnson, „Audio EQ Cookbook” (RBJ biquad formulas)", "https://www.w3.org/TR/audio-eq-cookbook/", "14.06.2025"),
        ("M. R. Schroeder, „Natural Sounding Artificial Reverberation”, J. Audio Eng. Soc., 1962; Jezar, Freeverb", "https://ccrma.stanford.edu/~jos/pasp/Freeverb.html", "14.06.2025"),
        ("Chris Wilson, „A Tale of Two Clocks – Scheduling Web Audio with Precision”", "https://web.dev/articles/audio-scheduling", "14.06.2025"),
        ("ITU-R BS.1770-4 – Algorithms to measure audio programme loudness", "https://www.itu.int/rec/R-REC-BS.1770", "14.06.2025"),
        ("W3C – AudioWorklet, Web Audio API specification", "https://www.w3.org/TR/webaudio/", "14.06.2025"),
        ("rustwasm – wasm-pack and the Rust and WebAssembly Book", "https://rustwasm.github.io/docs/book/", "15.06.2025"),
        ("MDN Web Docs – MediaDevices.enumerateDevices() și MediaRecorder API", "https://developer.mozilla.org/en-US/docs/Web/API/MediaDevices", "15.06.2025"),
    ]
    for i, (title, url, date) in enumerate(refs, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"[{i}] {title}"); r.font.name='Times New Roman'; r.font.size=Pt(11)
        p2 = doc.add_paragraph(); p2.paragraph_format.space_after = Pt(0); p2.paragraph_format.left_indent=Cm(0.6)
        r2 = p2.add_run(f"*** {url}"); r2.font.name='Times New Roman'; r2.font.size=Pt(10)
        p3 = doc.add_paragraph(); p3.paragraph_format.space_after = Pt(8); p3.paragraph_format.left_indent=Cm(0.6)
        r3 = p3.add_run(f"Accesat la {date}"); r3.font.name='Times New Roman'; r3.font.size=Pt(10); r3.italic=True

def back_matter(doc):
    # Lista figurilor
    h1(doc, "", "Lista figurilor și a tabelelor")
    body(doc, "În tabelul de mai jos sunt enumerate figurile și tabelele din lucrare, în ordinea apariției lor.", first_indent=False)
    pf = doc.add_paragraph(); rf = pf.add_run("Figuri"); rf.bold=True; rf.font.size=Pt(12); rf.font.name='Times New Roman'
    for n, cap in fig_index:
        p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
        r = p.add_run(f"Fig. {n} – {cap}"); r.font.name='Times New Roman'; r.font.size=Pt(11)
    pt = doc.add_paragraph(); rt = pt.add_run("Tabele"); rt.bold=True; rt.font.size=Pt(12); rt.font.name='Times New Roman'
    for n, cap in tab_index:
        p = doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
        r = p.add_run(f"Tabel {n} – {cap}"); r.font.name='Times New Roman'; r.font.size=Pt(11)
    # Declaratie
    doc.add_page_break()
    hp = doc.add_heading(level=1)
    r = hp.add_run("DECLARAȚIE DE AUTENTICITATE"); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True; r.font.color.rgb=DARKGREY
    body(doc, "Subsemnatul Tristan Cvasniuc, candidat la examenul de finalizare a studiilor de licență la Facultatea de Automatică și Calculatoare, specializarea Informatică, declar pe propria răspundere că lucrarea de față, intitulată „ResoLab – mini-DAW educațional pentru procesarea audio în timp real în browser folosind Rust și WebAssembly”, a fost elaborată de mine, pe baza propriei activități, și că nu conține porțiuni preluate din alte lucrări fără indicarea sursei.")
    body(doc, "Declar, de asemenea, că toate sursele utilizate, inclusiv documentația tehnică și articolele de specialitate, sunt citate corespunzător în text și menționate în bibliografie, iar fragmentele de cod prezentate provin din implementarea proprie a aplicației descrise.")
    p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(36)
    r = p.add_run("Data: ____________________            Semnătura: ____________________")
    r.font.name='Times New Roman'; r.font.size=Pt(12)
