# -*- coding: utf-8 -*-
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build(doc, h):
    body=h['body']; bullet=h['bullet']; H1=h['H1']; H2=h['H2']; H3=h['H3']
    fig_img=h['fig_img']; fig_ph=h['fig_ph']; code=h['code']; table_cap=h['table_cap']
    figlist=h['figlist']; tablist=h['tablist']

    # ============ 1. INTRODUCERE ============
    H1("1","Introducere")
    body("Acest capitol introductiv așază tema lucrării în contextul ei mai larg, explică motivele care au condus la alegerea ei și enunță obiectivele urmărite. În final este descrisă pe scurt organizarea întregii lucrări, pentru ca cititorul să poată parcurge fiecare secțiune cu o imagine clară asupra rolului pe care îl ocupă.")
    H2("1.1","Context și motivație")
    body("Prelucrarea numerică a sunetului a încetat de mult să mai fie apanajul studiourilor profesionale. Telefonul din buzunar înregistrează, montează și publică material audio, iar platformele de distribuție au coborât bariera de intrare aproape la zero. În acest peisaj, decalajul nu mai ține de accesul la unelte, ci de înțelegere: mulți utilizatori manevrează un compresor sau un egalizator fără a putea explica ce transformă acestea în semnal. Lucrarea se înscrie în zona aplicațiilor educaționale interactive pentru ingineria sunetului, la intersecția dintre prelucrarea numerică a semnalelor și tehnologiile web moderne.")
    body("Un program de tip stație de lucru audio digitală, cunoscut prin acronimul DAW, de la „digital audio workstation”, reunește înregistrarea, editarea, procesarea și mixarea sunetului într-un mediu unitar. Asemenea programe sunt, în mod tradițional, aplicații native voluminoase, care presupun instalare, licențiere și un calculator suficient de puternic. Apariția standardului WebAssembly a schimbat această ecuație, făcând posibilă rularea în browser a unui cod compilat, apropiat ca viteză de execuția nativă [1]. ResoLab valorifică tocmai această deschidere, mutând procesarea audio de performanță în pagina web.")
    body("Am ales această temă din dorința de a înțelege în profunzime atât prelucrarea de semnal, cât și tehnologiile care fac posibilă rularea ei în browser. Motivația principală a fost însă una pedagogică: instrumentele existente aplică efecte fără a le explica, lăsând utilizatorul fără un model mental al procesării. Proiectul nu își propune să concureze uneltele profesionale pe terenul numărului de funcții, ci să acopere o nișă distinctă — aceea a învățării asistate, în care utilizatorul vede, aude și citește, în același timp, ce face fiecare efect.")
    H2("1.2","Obiectivele lucrării")
    body("Pentru a delimita cu precizie ținta lucrării, am formulat de la început un set de obiective concrete, care au ghidat ulterior toate deciziile tehnice:")
    bullet("Construirea unui motor de procesare audio în timp real care să ruleze în browser cu latență redusă, scris în Rust și compilat către WebAssembly, capabil să susțină un lanț reconfigurabil de efecte fără întreruperi audibile.")
    bullet("Proiectarea unei interfețe care nu doar aplică efecte, ci explică funcționarea lor, prin explicații contextuale, feedback dinamic asupra parametrilor și vizualizări sincronizate cu redarea, adresate deopotrivă începătorilor și utilizatorilor avansați.")
    bullet("Implementarea unui set reprezentativ de procesoare audio pornind de la algoritmi consacrați în literatura de specialitate, astfel încât comportamentul lor să fie corect din punctul de vedere al ingineriei sunetului.")
    bullet("Asimilarea, în regim autodidact, a unui ansamblu de tehnologii noi pentru autor — limbajul Rust, modelul de execuție WebAssembly și interfața de programare Web Audio — și integrarea lor coerentă într-un produs funcțional.")
    bullet("Asigurarea unei experiențe complete de la capăt la capăt: încărcarea sunetului, procesarea lui, vizualizarea rezultatului și exportul într-un fișier WAV valid, redabil în orice player extern.")
    H2("1.3","Structura lucrării")
    body("Lucrarea este organizată în șase capitole. Capitolul de față introduce tema, motivația și obiectivele. Al doilea capitol analizează soluțiile similare existente și le compară sistematic cu propunerea de față. Al treilea capitol prezintă fundamentele teoretice și tehnologiile pe care se sprijină dezvoltarea. Al patrulea capitol, cel mai amplu, detaliază elaborarea proiectului, de la arhitectură până la implementarea fiecărei funcționalități. Al cincilea capitol descrie produsul finit și modul său de utilizare, iar ultimul capitol formulează concluziile și conturează direcțiile de dezvoltare ulterioară.")

    # ============ 2. SOLUTII SIMILARE ============
    H1("2","Soluții similare")
    body("Înainte de a justifica oportunitatea unei noi aplicații, este firesc să examinăm ce oferă deja peisajul instrumentelor existente. În acest capitol sunt analizate patru soluții reprezentative pentru editarea și producția audio în browser sau pe desktop, urmărind cu precădere funcționalitățile pe care le pun la dispoziție și tehnologiile pe care se sprijină. Capitolul se încheie cu o comparație sistematică, sintetizată într-un tabel.")
    H2("2.1","Audacity și portarea Wavacity")
    body("Audacity este probabil cel mai cunoscut editor audio gratuit și cu sursă deschisă, folosit de peste două decenii pentru înregistrare și editare multi-pistă [2]. Programul oferă un arsenal bogat de efecte, suport pentru numeroase formate și o comunitate vastă. Limitarea sa, din perspectiva temei de față, este de natură pedagogică: efectele sunt expuse sub formă de ferestre de dialog cu parametri numerici, fără a explica utilizatorului ce transformă fiecare reglaj asupra semnalului. Proiectul Wavacity reprezintă o portare a Audacity în browser, realizată prin compilarea codului original către WebAssembly [3]. Această inițiativă demonstrează concret că un editor audio complex poate rula într-o pagină web, validând premisa tehnologică pe care se sprijină și ResoLab, însă moștenește integral modelul de interacțiune al aplicației originale, inclusiv absența unui strat explicativ.")
    fig_ph("Interfața de editare a aplicației Audacity, cu fereastra de dialog a unui efect")
    H2("2.2","Soundtrap")
    body("Soundtrap, parte a ecosistemului Spotify, este o stație de lucru audio colaborativă care funcționează integral în browser [4]. Aplicația permite înregistrarea, aranjarea pe piste și aplicarea de efecte, punând accent pe colaborarea în timp real și pe integrarea cu serviciile educaționale. Diferența esențială ține de natura procesării: Soundtrap se bazează pe infrastructura cloud și pe un model de abonament, iar efectele sunt prezentate ca unelte gata făcute, fără a deschide o fereastră asupra mecanismelor interne. Componenta educațională vizează compoziția și producția muzicală, nu înțelegerea procesării de semnal la nivel de parametru.")
    H2("2.3","BandLab")
    body("BandLab este o platformă gratuită de creație muzicală, disponibilă în browser și pe dispozitive mobile, care combină o stație de lucru audio cu o rețea socială dedicată muzicienilor [5]. Modelul gratuit și accesibilitatea au făcut din ea o opțiune populară printre creatorii începători. Și în acest caz, orientarea este către rezultatul muzical, nu către procesul de învățare a ingineriei sunetului; efectele sunt aplicate prin interfețe simplificate, optimizate pentru rapiditate, iar procesarea se sprijină în bună măsură pe servere, ceea ce presupune o conexiune permanentă.")
    H2("2.4","AudioMass")
    body("AudioMass este un editor audio gratuit, cu sursă deschisă, care rulează în browser fără instalare sau cont [6]. Spre deosebire de soluțiile anterioare, este construit integral pe interfața de programare Web Audio nativă a browserului, fără un motor compilat separat. Această cale este simplă, însă oferă un control limitat asupra algoritmilor și nu permite implementarea propriilor procesoare cu precizia dorită. Lipsa unui strat educațional o menține, de asemenea, în categoria uneltelor de editare.")
    H2("2.5","Analiză comparativă")
    body("Punând alături soluțiile analizate, se conturează un tipar clar: instrumentele existente sunt fie editoare orientate către producție, fie platforme de creație muzicală, niciuna nefiind proiectată în jurul învățării procesării de semnal. Tabelul de mai jos sintetizează prezența sau absența unor caracteristici relevante pentru tema lucrării.")
    table_cap("Comparație între aplicațiile analizate și ResoLab")
    rows=[["Caracteristică","Audacity /\nWavacity","Soundtrap","BandLab","AudioMass","ResoLab"],
        ["Rulează în browser","Parțial","Da","Da","Da","Da"],
        ["Procesare locală (fără server)","Da","Nu","Nu","Da","Da"],
        ["Motor DSP propriu (Rust/WASM)","Da (C++)","Nu","Nu","Nu","Da"],
        ["Strat educațional / explicații","Nu","Parțial","Nu","Nu","Da"],
        ["Feedback dinamic pe parametri","Nu","Nu","Nu","Nu","Da"],
        ["Recomandări pe analiză spectrală","Nu","Nu","Nu","Nu","Da"],
        ["Vizualizări în timp real","Parțial","Parțial","Parțial","Da","Da"],
        ["Sintetizator integrat","Nu","Da","Da","Nu","Da"],
        ["Fără cont / fără abonament","Da","Nu","Parțial","Da","Da"]]
    t=doc.add_table(rows=len(rows),cols=6); t.style='Table Grid'; t.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for i,row in enumerate(rows):
        for j,val in enumerate(row):
            cell=t.rows[i].cells[j]; cell.text=''
            pp=cell.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.LEFT if j==0 else WD_ALIGN_PARAGRAPH.CENTER
            rr=pp.add_run(val); rr.font.name='Arial'; rr.font.size=Pt(9); rr.bold=(i==0 or j==5)
    body("Din tabel reiese că niciuna dintre aplicațiile consacrate nu îmbină procesarea locală de performanță cu un strat explicit de învățare. Soluțiile care rulează local nu explică procesarea, iar cele care includ resurse educaționale o fac la nivelul compoziției și depind de infrastructura cloud. ResoLab ocupă tocmai acest spațiu liber, ceea ce constituie principala sa contribuție față de stadiul actual.")

    # ============ 3. FUNDAMENTE TEORETICE ============
    H1("3","Fundamente teoretice")
    body("Acest capitol trece în revistă tehnologiile și bibliotecile pe care se sprijină dezvoltarea aplicației, oferind contextul necesar pentru înțelegerea deciziilor descrise ulterior. Prezentarea se concentrează asupra elementelor cu adevărat relevante pentru proiect, fără a epuiza documentația fiecărui instrument.")
    H2("3.1","Rust și WebAssembly")
    body("Rust este un limbaj de programare de sistem orientat către siguranță și performanță, care garantează absența unei categorii întregi de erori de memorie fără a recurge la un colector de gunoaie [7]. Tocmai această combinație îl face potrivit pentru cod audio: procesarea de semnal cere viteză predictibilă, iar pauzele introduse de un colector de gunoaie ar provoca întreruperi audibile. WebAssembly este un format binar de instrucțiuni de nivel scăzut, conceput ca țintă de compilare portabilă pentru limbaje precum C, C++ și Rust, executabil în browser la o viteză apropiată de cea nativă [1]. Codul Rust al motorului este compilat către acest format cu ajutorul utilitarului wasm-pack, iar comunicarea la granița dintre cele două lumi se realizează prin funcții exportate cu convenția C și prin schimb de pointeri către memoria liniară a modulului, evitând copierea inutilă a tampoanelor de eșantioane [8].")
    H2("3.2","Web Audio și AudioWorklet")
    body("Interfața Web Audio oferă browserului un graf de procesare a sunetului, alcătuit din noduri interconectate, prin care semnalul curge de la sursă către ieșire [9]. Pentru procesarea personalizată, standardul pune la dispoziție mecanismul AudioWorklet, care permite execuția de cod propriu pe firul dedicat audio, separat de firul principal al interfeței. Acest fir rulează cu prioritate ridicată, ceea ce reduce riscul întreruperilor. În arhitectura aplicației, procesorul AudioWorklet îndeplinește rolul unei punți subțiri: încarcă modulul WebAssembly, copiază tampoanele de intrare și ieșire și transmite mesajele de control. Procesarea propriu-zisă se desfășoară în blocuri de o sută douăzeci și opt de eșantioane, dimensiunea standard a cuantei de redare din Web Audio.")
    H2("3.3","React și TypeScript")
    body("Interfața grafică este construită cu React, o bibliotecă pentru construirea de interfețe pe bază de componente, în care reprezentarea vizuală este o funcție a stării [10]. Modelul declarativ se potrivește bine unei aplicații cu multe controale interdependente. Peste React am folosit TypeScript, o extensie a limbajului JavaScript care adaugă un sistem de tipuri verificat la compilare [11]. Tipizarea statică s-a dovedit deosebit de utilă într-un proiect cu numeroase mesaje schimbate între interfață și motorul audio, deoarece a permis surprinderea timpurie a nepotrivirilor de structură.")
    H2("3.4","Gestiunea stării cu Zustand")
    body("Pentru coordonarea stării globale am ales biblioteca Zustand, o soluție minimalistă care expune starea prin depozite independente, accesate cu ajutorul unor selectoare [12]. În locul unui depozit unic și monolitic, starea este împărțită pe domenii: lanțul de efecte, fișierul audio și redarea, datele de analiză, starea educațională, preseturile, interfața și sintetizatorul. Fiecare componentă se abonează doar la felia de stare de care are nevoie, ceea ce limitează redesenările inutile, iar logica de modificare a stării este definită în interiorul depozitelor, sub formă de acțiuni.")
    H2("3.5","Biblioteci pentru interfață și vizualizare")
    body("Aspectul vizual se sprijină pe TailwindCSS, un sistem de stilizare bazat pe clase utilitare, completat de primitivele accesibile oferite de Radix UI [13]. Reprezentarea formei de undă este realizată cu biblioteca wavesurfer.js, care desenează semnalul și gestionează selecția regiunilor de buclă [14]. Pentru afișarea spectrului în timp real am folosit audiomotion-analyzer, iar persistența preseturilor se realizează printr-o bază de date IndexedDB, accesată prin biblioteca idb, în timp ce preferințele de interfață sunt salvate în spațiul local de stocare al browserului.")
    print("chapters 1-3 built")

    # ============ 4. ELABORAREA PROIECTULUI ============
    H1("4","Elaborarea proiectului")
    body("Capitolul de față coboară la nivelul detaliilor concrete, prezentând modul în care a fost gândit și construit sistemul, de la organizarea generală a componentelor până la implementarea fiecărei funcționalități. Sunt descrise aici deciziile structurale majore, traseul semnalului, motorul de procesare și realizarea propriu-zisă a principalelor module, însoțite de fragmente de cod relevante.")
    H2("4.1","Arhitectura generală")
    body("Aplicația este împărțită în trei straturi care comunică pe căi bine delimitate. Primul strat este interfața grafică, scrisă în React și TypeScript, responsabilă de afișare și de captarea acțiunilor utilizatorului. Al doilea strat este puntea audio, materializată într-un procesor AudioWorklet care rulează pe firul dedicat sunetului. Al treilea strat este motorul de procesare, scris în Rust și compilat către WebAssembly, care efectuează toate calculele de semnal. Această separare reflectă o regulă strictă respectată în tot proiectul: nicio prelucrare audio nu se execută în JavaScript, ci numai în codul nativ. Organizarea pe straturi este ilustrată în figura următoare.")
    fig_img("fig_architecture.png","Arhitectura pe trei straturi a aplicației ResoLab", 13)
    body("Comunicarea dintre straturi urmează un flux unidirecțional, ușor de urmărit. Atunci când utilizatorul mișcă un buton rotativ, valoarea ajunge mai întâi în depozitul de stare corespunzător, de unde este transmisă prin mesaj către procesorul AudioWorklet, care la rândul lui apelează funcția potrivită din motorul nativ. În sens invers, datele de analiză sunt colectate de motor, trimise periodic către firul principal și depuse într-un depozit de analiză, de unde alimentează vizualizările. Această circulație clară a informației, redată în figura de mai jos, a simplificat considerabil depanarea.")
    fig_img("fig_dataflow.png","Circulația datelor: fluxul de control și fluxul de analiză", 15)
    H2("4.2","Traseul semnalului audio")
    body("Pentru a înțelege funcționarea sistemului, este util să urmărim drumul complet al semnalului. Un fișier audio este mai întâi decodat local, prin mecanismul oferit de browser, rezultând un tampon de eșantioane în virgulă mobilă. La redare, acest tampon este furnizat printr-un nod sursă către procesorul AudioWorklet, care îl trece, bloc cu bloc, prin motorul nativ. Rezultatul ajunge la difuzoarele sau căștile utilizatorului, iar în paralel un nod de analiză derivă o ramură a semnalului către vizualizări, fără a altera traseul principal. Atât sintetizatorul, cât și intrarea de microfon se injectează în același punct, înaintea lanțului de efecte, astfel încât prelucrarea ulterioară este identică indiferent de sursă.")
    fig_img("fig_signalflow.png","Traseul semnalului audio prin sistem", 15)
    H2("4.3","Motorul de efecte stereo")
    body("Inima sistemului este un motor stereo construit pe un principiu pe care l-am numit, în cod, proiectarea cu instanțe duble. Fiecare poziție din lanțul de efecte nu conține un singur procesor, ci două instanțe independente ale aceluiași efect — una pentru canalul stâng și una pentru cel drept — ai căror parametri sunt menținuți sincronizați. Am ales această abordare deoarece păstrează intactă imaginea stereo a semnalului prin întregul lanț, fără a modifica implementarea fiecărui efect. Principiul este redat în figura următoare.")
    fig_img("fig_stereo.png","Motorul stereo cu instanțe duble pe canal", 14)
    body("Structura de date care reprezintă o poziție din lanț reține un identificator, starea de ocolire și cele două instanțe ale efectului, reunite sub o interfață comună. Toate efectele respectă acest tip-trăsătură, ceea ce permite tratarea lor uniformă, indiferent de algoritmul intern.")
    code("struct EffectSlot {\n"
         "    id: u32,\n"
         "    bypassed: bool,\n"
         "    left: Box<dyn Effect>,\n"
         "    right: Box<dyn Effect>,\n"
         "}\n\n"
         "pub struct Engine {\n"
         "    sample_rate: f32,\n"
         "    chain: Vec<EffectSlot>,\n"
         "    work_l: Vec<f32>, work_r: Vec<f32>,\n"
         "    temp_l: Vec<f32>, temp_r: Vec<f32>,\n"
         "}",
         "Structura motorului stereo și a unei poziții din lanțul de efecte")
    body("Tampoanele de lucru sunt preluate alternativ, după modelul ping-pong, astfel încât ieșirea unui efect devine intrarea următorului, fără copieri suplimentare. Toate tampoanele sunt alocate o singură dată, la inițializare; pe traseul critic de procesare nu se mai face nicio alocare de memorie, întrucât o alocare neașteptată pe firul audio s-ar traduce într-un artefact sonor.")
    H2("4.4","Procesoarele audio implementate")
    body("Aplicația include șaisprezece procesoare audio, fiecare implementat ca un modul Rust independent. Ele acoperă principalele categorii de prelucrare: dinamică, frecvențială, temporală și de saturație. În continuare sunt detaliate cele mai reprezentative, alese pentru a ilustra abordările diferite.")
    H3("4.4.1","Compresorul dinamic")
    body("Compresorul reduce diferența dintre pasajele tari și cele slabe ale semnalului, fiind unul dintre cele mai folosite, dar și cele mai greu de înțeles efecte. Implementarea urmează modelul feed-forward descris în literatura de specialitate, cu un genunchi moale care netezește tranziția în jurul pragului [15]. Semnalul de comandă trece printr-un detector de anvelopă, după care o funcție de câștig calculează reducerea necesară.")
    code("pub const PARAM_THRESHOLD_DB: u32 = 0;\n"
         "pub const PARAM_RATIO: u32 = 1;\n"
         "pub const PARAM_ATTACK_MS: u32 = 2;\n"
         "pub const PARAM_RELEASE_MS: u32 = 3;\n"
         "pub const PARAM_KNEE_DB: u32 = 4;\n"
         "pub const PARAM_MAKEUP_DB: u32 = 5;\n"
         "// detector de anvelopă -> calcul câștig -> makeup -> mix uscat/umed",
         "Parametrii compresorului și ordinea etapelor de procesare")
    body("Valoarea reducerii de câștig este expusă către interfață pentru a alimenta un indicator vizual dedicat, astfel încât utilizatorul să poată vedea, nu doar auzi, cât de mult intervine compresorul. Toți parametrii sunt netezi în timp printr-un filtru de ordinul întâi, ceea ce elimină zgomotul de tip fermoar care ar apărea la modificări bruște.")
    fig_ph("Cardul compresorului în interfață, cu indicatorul de reducere a câștigului")
    H3("4.4.2","Egalizatorul parametric")
    body("Egalizatorul parametric ajustează balanța frecvențială a semnalului prin patru benzi în cascadă, fiecare putând lua forma unui filtru de tip clopot, prag jos sau înalt, trecere sus sau jos, ori filtru de tăiere. Coeficienții fiecărei benzi sunt recalculați doar la modificarea parametrilor, după formulele consacrate ale lui Robert Bristow-Johnson [16], iar procesarea propriu-zisă constă într-o cascadă de filtre biquad. Pe lângă procesarea în Rust, interfața afișează o curbă de răspuns în frecvență, calculată separat pentru reprezentare vizuală, care permite utilizatorului să anticipeze efectul reglajelor înainte de a le asculta.")
    fig_ph("Curba de răspuns în frecvență a egalizatorului parametric cu patru benzi")
    H3("4.4.3","Reverberația")
    body("Reverberația simulează reflexiile sunetului într-un spațiu, conferind profunzime semnalului. Implementarea se bazează pe arhitectura Schroeder, cu patru filtre piepten cu reacție și atenuare dispuse în paralel, urmate de două filtre trece-tot în serie [17]. Lungimile liniilor de întârziere pornesc de la valorile clasice ale algoritmului Freeverb, scalate la frecvența de eșantionare curentă.")
    code("// patru filtre piepten în paralel + două filtre trece-tot în serie\n"
         "const COMB_LENGTHS_44K: [usize; 4] = [1116, 1188, 1277, 1356];\n"
         "const ALLPASS_LENGTHS_44K: [usize; 2] = [556, 441];\n"
         "// atenuarea în bucla de reacție este un filtru trece-jos de ordinul întâi",
         "Configurația liniilor de întârziere ale reverberației de tip Schroeder")
    body("Parametrul de atenuare controlează un filtru trece-jos plasat în bucla de reacție a fiecărui filtru piepten, ceea ce face ca frecvențele înalte să se stingă mai repede decât cele joase, întocmai ca într-un spațiu real. Un predelay, realizat cu o linie de întârziere separată, permite distanțarea reverberației de sunetul direct.")
    H3("4.4.4","Celelalte procesoare")
    body("Pe lângă cele descrise, motorul include un câștig simplu, o poartă de zgomot, un limitator, o întârziere cu reacție, saturație, cor, flanger, deplasare de ton, phaser, modelator de tranzienți, de-esser, expandor și un reductor de zgomot. Fiecare respectă aceeași interfață comună și aceeași disciplină a alocării, diferind doar prin algoritmul intern. Această uniformitate a permis adăugarea treptată de noi efecte fără a modifica motorul.")
    H2("4.5","Puntea AudioWorklet")
    body("Procesorul AudioWorklet face legătura între interfață și motorul nativ, fără a efectua el însuși vreo prelucrare de semnal. La inițializare, primește octeții modulului WebAssembly, îl instanțiază și alocă tampoanele de intrare și ieșire pentru ambele canale. Mesajele de control sosite de la firul principal sunt traduse în apeluri către funcțiile motorului.")
    code("process(inputs, outputs) {\n"
         "  if (!this.ready) return true;\n"
         "  this.inputView.set(input[0]);\n"
         "  this.inputRView.set(input[1] || input[0]);\n"
         "  this.wasm.process_stereo(this.enginePtr,\n"
         "     this.inputPtr, this.inputRPtr,\n"
         "     this.outputPtr, this.outputRPtr, RENDER_QUANTUM);\n"
         "  output[0].set(this.outputView);\n"
         "  output[1].set(this.outputRView);\n"
         "  return true;\n"
         "}",
         "Bucla de procesare a punții AudioWorklet")
    body("Pentru a evita copierea eșantioanelor, codul folosește vederi de tip tablou tipizat direct peste memoria liniară a modulului WebAssembly. Atunci când motorul detectează o ocolire globală, semnalul este transmis nemodificat de la intrare la ieșire, ceea ce stă la baza comparației rapide între semnalul procesat și cel original.")
    H2("4.6","Sintetizatorul")
    body("Sintetizatorul este un instrument substractiv polifonic, capabil să redea simultan opt voci. Fiecare voce dispune de două oscilatoare cu posibilitate de dezacordare, o anvelopă de tip atac–cădere–susținere–revenire, un filtru variabil de stare și un oscilator de joasă frecvență pentru modulație. Vocile sunt alocate după vârstă, cea mai veche fiind sacrificată atunci când toate sunt ocupate. Comanda poate veni fie de la tastatura calculatorului, fie de la un dispozitiv MIDI conectat prin interfața Web MIDI. Peste mecanismul de bază am adăugat un mod de acorduri și un arpegiator, iar sunetul rezultat se injectează în lanțul de efecte, putând fi modelat suplimentar.")
    fig_ph("Panoul sintetizatorului, cu oscilatoarele, filtrul și anvelopa")
    H2("4.7","Înregistrarea și selecția interfeței audio")
    body("Modulul de înregistrare captează ieșirea procesată a aplicației și o salvează într-un fișier audio. Pentru a oferi flexibilitate utilizatorilor cu echipament profesional, butonul de microfon deschide un selector care enumeră dispozitivele de intrare disponibile. Dacă browserul nu a primit încă permisiunea de acces, etichetele dispozitivelor apar goale, motiv pentru care codul solicită mai întâi permisiunea, apoi reia enumerarea.")
    code("let devices = await navigator.mediaDevices.enumerateDevices()\n"
         "let inputs = devices.filter((d) => d.kind === 'audioinput')\n"
         "if (inputs.length > 0 && inputs[0].label === '') {\n"
         "  const tmp = await navigator.mediaDevices.getUserMedia({ audio: true })\n"
         "  tmp.getTracks().forEach((t) => t.stop())\n"
         "  devices = await navigator.mediaDevices.enumerateDevices()\n"
         "  inputs = devices.filter((d) => d.kind === 'audioinput')\n"
         "}",
         "Enumerarea dispozitivelor de intrare cu solicitarea prealabilă a permisiunii")
    body("După alegerea unui dispozitiv, semnalul este captat de la acea sursă și conectat în motor. Un comutator de monitorizare permite ascultarea directă a intrării, însoțit de un avertisment privind riscul de microfonie atunci când nu se folosesc căști. Toate erorile posibile — permisiune refuzată, dispozitiv indisponibil — sunt afișate în interfață, nu doar consemnate în consolă. O atenție deosebită a necesitat cronometrul de înregistrare, care inițial rămânea blocat la zero; soluția a constat în declanșarea unui interval ce incrementează un contor de secunde, oprit la finalul înregistrării, cu o gardă împotriva pornirii duble.")
    fig_ph("Selectorul de interfață audio și comutatorul de monitorizare")
    H2("4.8","Metronomul cu programare anticipată")
    body("Metronomul produce un click ritmic stabil, util pentru exersare. Provocarea principală a fost precizia temporală: un cronometru obișnuit din JavaScript suferă de fluctuații care ar face pulsul neregulat. Soluția adoptată este tehnica programării anticipate, recunoscută ca standard în domeniul Web Audio [18]. Un cronometru rar inspectează periodic o fereastră de timp apropiată și programează în avans toate bătăile care urmează să cadă în interiorul ei, conform schemei din figura următoare.")
    fig_img("fig_scheduler.png","Programarea anticipată a bătăilor metronomului", 14)
    code("function scheduler() {\n"
         "  while (nextNoteTime < ctx.currentTime + LOOKAHEAD_SEC) {\n"
         "    const accented = currentBeat === 0\n"
         "    scheduleClick(nextNoteTime, currentBeat, accented)\n"
         "    nextNoteTime += 60.0 / bpm\n"
         "    currentBeat = (currentBeat + 1) % beatsPerMeasure\n"
         "  }\n"
         "  timerID = setTimeout(scheduler, SCHEDULE_INTERVAL_MS)\n"
         "}",
         "Programatorul anticipat al metronomului")
    body("Fiecare click este generat de un oscilator de scurtă durată, fără a recurge la fișiere audio externe. Prima bătaie a fiecărei măsuri folosește o frecvență mai înaltă și un volum mai mare, marcând accentul, în funcție de metrul ales. Tempoul poate fi reglat între patruzeci și două sute patruzeci de bătăi pe minut, iar modificarea sa în timpul rulării actualizează imediat programarea, fără a opri pulsul.")
    H2("4.9","Exportul în format WAV")
    body("La finalul prelucrării, utilizatorul poate exporta rezultatul într-un fișier WAV. Exportul nu se bazează pe redarea în timp real, ci pe o procesare în regim offline: tamponul sursă este trecut, bloc cu bloc, printr-o instanță proaspătă a motorului, cu același lanț de efecte, după o etapă de încălzire care lasă filtrele și netezitoarele să se stabilizeze. Un detaliu subtil, dar important, a fost corectitudinea antetului WAV: la scrierea în flux, dimensiunile finale nu sunt cunoscute de la început, ceea ce poate produce un fișier a cărui durată este afișată greșit. Soluția a fost rescrierea antetului cu valorile corecte ale dimensiunii totale și ale blocului de date după ce toate eșantioanele au fost adunate.")
    H2("4.10","Stratul educațional")
    body("Componenta care diferențiază cel mai puternic ResoLab de instrumentele existente este stratul educațional. Acesta nu se rezumă la texte statice, ci analizează configurația curentă a lanțului de efecte și formulează observații adaptate valorilor reale ale parametrilor. Un motor de reguli examinează fiecare efect și semnalează situațiile problematice — un prag prea coborât, un raport exagerat, o ordonare nepotrivită — oferind explicații în limbaj accesibil sau tehnic. În completare, un motor de recomandări analizează conținutul spectral al semnalului și propune efecte potrivite, pe baza unor reguli euristice. Întreaga logică se bazează pe analiză de semnal, nu pe metode de învățare automată, ceea ce o face transparentă și explicabilă — o calitate aliniată cu scopul didactic al aplicației.")
    fig_ph("Panoul de feedback dinamic și panoul de recomandări")
    print("chapter 4 built")

    # ============ 5. DESCRIEREA SI UTILIZAREA PRODUSULUI ============
    H1("5","Descrierea și utilizarea produsului")
    body("Acest capitol prezintă aplicația din perspectiva utilizatorului final, urmărind pașii pe care îi parcurge de la primul contact cu interfața până la obținerea unui rezultat exportat. Descrierea însoțește principalele ecrane și acțiuni, oferind o imagine de ansamblu asupra modului de operare.")
    H2("5.1","Privire de ansamblu asupra interfeței")
    body("Interfața este organizată ca un atelier cu trei coloane. În stânga se află un panou de navigare care găzduiește eșantioanele demonstrative și preseturile, în centru se desfășoară forma de undă și lanțul de efecte, iar în dreapta se află panoul de inspecție, cu vizualizările și stratul educațional. În partea de jos, o bară de transport reunește comenzile de redare, controlul de tempo, metronomul și înregistrarea. Această dispunere a fost gândită pentru a păstra permanent în câmpul vizual atât semnalul, cât și efectul fiecărei acțiuni.")
    fig_ph("Dispunerea generală a interfeței ResoLab, cu cele trei coloane și bara de transport")
    H2("5.2","Încărcarea fișierului și redarea")
    body("La deschiderea aplicației, utilizatorul întâmpină un spațiu central pregătit să primească un fișier audio. Acesta poate fi adus prin tragere și plasare sau selectat cu un clic, iar pentru explorare rapidă sunt disponibile câteva eșantioane demonstrative în panoul din stânga. Odată încărcat, fișierul este afișat sub forma undei sonore, iar bara de transport devine activă. Pe forma de undă, utilizatorul poate fixa o regiune de buclă prin tragere și poate sări într-un punct anume printr-un clic, interacțiuni dublate de scurtături de tastatură.")
    fig_ph("Ecranul inițial cu zona de încărcare a fișierului audio")
    H2("5.3","Construirea lanțului de efecte")
    body("Adăugarea unui efect se face din butonul dedicat, care deschide lista procesoarelor disponibile. Efectele se așază în ordine, formând un lanț pe care utilizatorul îl poate rearanja prin tragere, folosind mânerul afișat pe fiecare card. Fiecare efect poate fi dezactivat individual, iar întregul lanț poate fi ocolit dintr-o singură comandă, pentru a compara rapid semnalul procesat cu cel original. La trecerea cursorului peste un control apare o explicație contextuală, iar panoul din dreapta comentează configurația. Orice modificare poate fi anulată sau refăcută, ceea ce încurajează experimentarea.")
    fig_ph("Lanțul de efecte cu mai multe procesoare active")
    H2("5.4","Vizualizări și feedback")
    body("Panoul de vizualizare oferă patru perspective complementare: spectrul instantaneu, spectrograma care urmărește evoluția spectrului în timp, osciloscopul care arată forma de undă și un mod de comparație care îngheață un spectru de referință alături de semnalul curent. Pe lângă acestea, interfața afișează indicatoare de nivel pentru vârf și valoare eficace, un indicator de reducere a câștigului și o măsură a tăriei sonore integrate, calculată după recomandarea internațională în domeniu [19]. Aceste instrumente oferă utilizatorului un limbaj vizual prin care poate corela ceea ce aude cu ceea ce se întâmplă cu semnalul.")
    fig_ph("Panoul de vizualizare cu spectrul în timp real și spectrograma")
    H2("5.5","Sintetizator, metronom și înregistrare")
    body("Dincolo de prelucrarea fișierelor, utilizatorul poate genera sunete de la zero cu ajutorul sintetizatorului, cântând de la tastatură sau de la un dispozitiv MIDI, și le poate trece prin lanțul de efecte. Metronomul îl asistă în exersarea ritmică, oferind reglaj de tempo, metru și un indicator vizual de puls. Cei care dispun de o interfață audio externă pot alege dispozitivul de intrare, pot asculta semnalul în timp real și îl pot înregistra direct în aplicație.")
    H2("5.6","Preseturi și export")
    body("Configurațiile reușite pot fi salvate ca preseturi personalizate, cu un nume ales de utilizator, păstrate local într-o bază de date a browserului, astfel încât rămân disponibile și la vizitele ulterioare. Pe lângă preseturile proprii, aplicația include un set de configurații de fabrică. Când rezultatul este satisfăcător, utilizatorul îl poate exporta într-un fișier WAV printr-o singură comandă; fișierul descărcat poate fi redat în orice player extern, cu durata și calitatea corecte, închizând astfel ciclul complet de lucru.")
    fig_ph("Exportul rezultatului procesat într-un fișier WAV")
    H2("5.7","Tutorialul ghidat")
    body("Pentru a-i orienta pe utilizatorii aflați la prima vizită, aplicația include un tutorial ghidat de tip reflector. Acesta întunecă întreaga interfață, decupând doar zona relevantă pentru pasul curent, în jurul căreia plasează o casetă explicativă. Poziția decupajului este calculată dinamic, pe baza dimensiunilor reale ale elementului vizat, iar caseta este așezată inteligent astfel încât să nu iasă din cadru. Tutorialul parcurge șase pași, de la încărcarea unui fișier până la export, reducând semnificativ pragul de intrare pentru utilizatorii fără experiență.")
    fig_ph("Un pas din tutorialul ghidat, cu zona evidențiată și caseta explicativă")

    # ============ 6. CONCLUZII ============
    H1("6","Concluzii")
    body("În urma parcurgerii întregului proces de proiectare și dezvoltare, capitolul de față sintetizează rezultatele obținute, problemele întâmpinate și modul în care au fost depășite, contribuțiile aplicației față de soluțiile existente și direcțiile pe care le considerăm promițătoare pentru viitor.")
    H2("6.1","Rezultate obținute")
    body("Obiectivele formulate la începutul lucrării au fost atinse. Am construit un motor de procesare audio care rulează în browser cu latență redusă, scris în Rust și compilat către WebAssembly, capabil să susțină un lanț reconfigurabil de șaisprezece efecte fără întreruperi audibile. Am realizat o interfață care nu doar aplică efecte, ci le explică, prin explicații contextuale, feedback dinamic și vizualizări sincronizate, adaptate atât începătorilor, cât și utilizatorilor avansați. Procesoarele audio au fost implementate pornind de la algoritmi consacrați, ceea ce le conferă un comportament corect din punctul de vedere al ingineriei sunetului. Nu în ultimul rând, am asigurat un ciclu complet de lucru, de la încărcarea sunetului până la exportul într-un fișier WAV valid.")
    H2("6.2","Probleme întâmpinate și soluțiile adoptate")
    body("Pe parcursul dezvoltării au apărut mai multe dificultăți tehnice. Comunicarea la granița dintre JavaScript și WebAssembly a impus o atenție deosebită gestiunii memoriei, întrucât creșterea memoriei modulului invalidează vederile către tampoane; soluția a fost reobținerea acestor vederi după fiecare operație care putea provoca o realocare. Precizia temporală a metronomului, afectată inițial de fluctuațiile cronometrelor obișnuite, a fost rezolvată prin tehnica programării anticipate. Durata greșită a fișierelor exportate, cauzată de un antet scris înainte de cunoașterea dimensiunii finale, a fost corectată prin rescrierea antetului cu valorile exacte. Fiecare dintre aceste obstacole a constituit, retrospectiv, o ocazie de aprofundare.")
    H2("6.3","Contribuții față de soluțiile existente")
    body("Comparativ cu aplicațiile analizate în al doilea capitol, ResoLab aduce o combinație pe care niciuna dintre ele nu o oferă: procesare locală de performanță, realizată cu un motor propriu, dublată de un strat educațional care comentează și recomandă în funcție de starea reală a semnalului. Acolo unde editoarele consacrate aplică efecte fără a le explica, iar platformele de creație se sprijină pe infrastructura cloud și vizează compoziția, aplicația de față pune înțelegerea procesării în centrul experienței. Această orientare didactică, susținută tehnic de un motor rapid și transparent, reprezintă contribuția principală a lucrării.")
    H2("6.4","Direcții de dezvoltare ulterioară")
    body("Aplicația rămâne deschisă mai multor direcții de extindere. Programarea temporală a metronomului și a arpegiatorului ar putea fi mutată pe o bază mai precisă, sincronizată cu ceasul audio. Suportul MIDI ar putea fi îmbogățit cu mesaje de control continuu și cu sensibilitate la viteză. Stratul educațional s-ar putea extinde cu lecții interactive structurate și cu exerciții ghidate. De asemenea, exportul fișierelor de mari dimensiuni ar beneficia de un indicator de progres și de o procesare eșalonată. Aceste direcții depășesc scopul lucrării de față, dar conturează un parcurs firesc de maturizare a proiectului.")

    # ============ BIBLIOGRAFIE ============
    H1("","Bibliografie")
    refs=[
    'WebAssembly, „WebAssembly Specification”, *** https://webassembly.org/, accesare iunie 2026.',
    'Audacity Team, „Audacity – Free, open source, cross-platform audio software”, *** https://www.audacityteam.org/, accesare iunie 2026.',
    'Wavacity, „Wavacity – Audacity ported to the browser via WebAssembly”, *** https://wavacity.com/, accesare iunie 2026.',
    'Spotify, „Soundtrap – Make music online”, *** https://www.soundtrap.com/, accesare iunie 2026.',
    'BandLab Technologies, „BandLab – Free music making platform”, *** https://www.bandlab.com/, accesare iunie 2026.',
    'P. Bernat, „AudioMass – Free, open source web audio editor”, *** https://audiomass.co/, accesare iunie 2026.',
    'S. Klabnik și C. Nichols, „The Rust Programming Language”, No Starch Press, 2023. *** https://doc.rust-lang.org/book/, accesare iunie 2026.',
    'Rust and WebAssembly Working Group, „The `wasm-pack` Book”, *** https://rustwasm.github.io/docs/wasm-pack/, accesare iunie 2026.',
    'Mozilla Developer Network, „Web Audio API”, *** https://developer.mozilla.org/en-US/docs/Web/API/Web_Audio_API, accesare iunie 2026.',
    'Meta Open Source, „React – Documentația oficială”, *** https://react.dev/, accesare iunie 2026.',
    'Microsoft, „TypeScript – Documentation”, *** https://www.typescriptlang.org/docs/, accesare iunie 2026.',
    'Poimandres, „Zustand – State management for React”, *** https://zustand.docs.pmnd.rs/, accesare iunie 2026.',
    'Tailwind Labs, „Tailwind CSS – Documentation”; WorkOS, „Radix UI Primitives”, *** https://tailwindcss.com/docs, accesare iunie 2026.',
    'wavesurfer.js, „wavesurfer.js – Audio waveform visualisation”, *** https://wavesurfer.xyz/, accesare iunie 2026.',
    'U. Zölzer (ed.), „DAFX: Digital Audio Effects”, ediția a 2-a, *Wiley*, 2011.',
    'R. Bristow-Johnson, „Cookbook formulae for audio EQ biquad filter coefficients”, *Audio EQ Cookbook*, W3C, 2021. *** https://www.w3.org/TR/audio-eq-cookbook/, accesare iunie 2026.',
    'M. R. Schroeder, „Natural Sounding Artificial Reverberation”, *Journal of the Audio Engineering Society*, vol. 10, nr. 3, 1962.',
    'C. Wilson, „A Tale of Two Clocks – Scheduling Web Audio with Precision”, *** https://web.dev/articles/audio-scheduling, accesare iunie 2026.',
    'Uniunea Internațională a Telecomunicațiilor, „Recommendation ITU-R BS.1770-4: Algorithms to measure audio programme loudness and true-peak audio level”, 2015. *** https://www.itu.int/rec/R-REC-BS.1770, accesare iunie 2026.',
    'W3C, „Web Audio API – W3C Recommendation”, *** https://www.w3.org/TR/webaudio/, accesare iunie 2026.',
    'Mozilla Developer Network, „MediaDevices.enumerateDevices() și MediaRecorder API”, *** https://developer.mozilla.org/en-US/docs/Web/API/MediaDevices, accesare iunie 2026.',
    ]
    for i,r in enumerate(refs,1):
        p=doc.add_paragraph(style='Normal'); p.paragraph_format.space_after=Pt(6)
        run=p.add_run(f"[{i}] {r}"); run.font.name='Arial'; run.font.size=Pt(11)

    # ============ LISTA FIGURILOR SI TABELELOR ============
    H1("","Lista figurilor și a tabelelor")
    body("În continuare sunt enumerate figurile și tabelele din lucrare, în ordinea apariției lor.", indent=False)
    pf=doc.add_paragraph(); rf=pf.add_run("Figuri"); rf.bold=True; rf.font.name='Arial'; rf.font.size=Pt(12)
    for n,cap in figlist:
        p=doc.add_paragraph(style='Normal'); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(f"Figura {n} – {cap}"); r.font.name='Arial'; r.font.size=Pt(11)
    pt=doc.add_paragraph(); rt=pt.add_run("Tabele"); rt.bold=True; rt.font.name='Arial'; rt.font.size=Pt(12)
    for n,cap in tablist:
        p=doc.add_paragraph(style='Normal'); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(f"Tabelul {n} – {cap}"); r.font.name='Arial'; r.font.size=Pt(11)

    # ============ DECLARATIE ============
    H1("","Declarație de originalitate")
    body("Subsemnatul " + h['AUTHOR'].replace('CVASNIUC','Cvasniuc') + ", candidat la examenul de finalizare a studiilor de licență la Facultatea de Automatică și Calculatoare a Universității Politehnica Timișoara, specializarea Informatică, declar pe propria răspundere că lucrarea de față, intitulată „ResoLab – mini-DAW educațional pentru procesarea audio în timp real în browser folosind Rust și WebAssembly”, a fost elaborată de mine, pe baza propriei activități, și că nu conține porțiuni preluate din alte lucrări fără indicarea sursei.")
    body("Declar, de asemenea, că toate sursele utilizate, inclusiv documentația tehnică și articolele de specialitate, sunt citate corespunzător în text și menționate în bibliografie, iar fragmentele de cod prezentate provin din implementarea proprie a aplicației descrise.")
    body("Notă: formularul oficial al declarației de originalitate, disponibil la adresa https://media.upt.ro/Declaratie_aut_studii.pdf, va fi completat, semnat cu pix albastru și atașat ca ultimă pagină a lucrării.", indent=False)
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(36)
    r=p.add_run("Data: ____________________            Semnătura: ____________________")
    r.font.name='Arial'; r.font.size=Pt(12)
    print("chapters 5-6 + back matter built")
